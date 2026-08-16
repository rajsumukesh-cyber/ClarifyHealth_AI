import os
import re
from typing import Tuple, List, Dict, Any
from pypdf import PdfReader
import docx
from PIL import Image
import io

class DocumentExtractionResult:
    def __init__(self, text: str, page_count: int, unclear_sections: List[str], metadata: Dict[str, Any]):
        self.text = text
        self.page_count = page_count
        self.unclear_sections = unclear_sections
        self.metadata = metadata

class FileExtractor:
    @staticmethod
    def extract_from_pdf(file_path: str) -> DocumentExtractionResult:
        """Extract text and metadata from PDF files."""
        extracted_pages = []
        unclear_sections = []
        try:
            reader = PdfReader(file_path)
            page_count = len(reader.pages)
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                page_text = page_text.strip()
                if not page_text:
                    unclear_sections.append(f"Page {idx + 1}: Image-only or non-selectable text detected.")
                else:
                    extracted_pages.append(f"--- Page {idx + 1} ---\n{page_text}")
            
            full_text = "\n\n".join(extracted_pages).strip()
            if not full_text:
                full_text = "[No machine-readable text found in PDF. This document may be scanned or image-based.]"
                unclear_sections.append("Entire document appears to be scanned or contains non-extractable text.")
            
            return DocumentExtractionResult(
                text=FileExtractor.clean_extracted_text(full_text),
                page_count=page_count,
                unclear_sections=unclear_sections,
                metadata={"format": "PDF", "num_pages": page_count}
            )
        except Exception as e:
            return DocumentExtractionResult(
                text=f"[Error reading PDF: {str(e)}]",
                page_count=1,
                unclear_sections=[f"Could not parse PDF structure: {str(e)}"],
                metadata={"error": str(e)}
            )

    @staticmethod
    def extract_from_docx(file_path: str) -> DocumentExtractionResult:
        """Extract text and tables from Word DOCX files."""
        try:
            doc = docx.Document(file_path)
            content_blocks = []
            
            # Paragraphs
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    content_blocks.append(text)
            
            # Tables
            for table_idx, table in enumerate(doc.tables):
                table_lines = []
                for row in table.rows:
                    row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    if any(row_cells):
                        table_lines.append(" | ".join(row_cells))
                if table_lines:
                    content_blocks.append(f"\n[Table {table_idx + 1}]\n" + "\n".join(table_lines))
            
            full_text = "\n\n".join(content_blocks).strip()
            if not full_text:
                full_text = "[Empty DOCX document]"
            
            return DocumentExtractionResult(
                text=FileExtractor.clean_extracted_text(full_text),
                page_count=max(1, len(full_text) // 2500 + 1),
                unclear_sections=[],
                metadata={"format": "DOCX", "tables_count": len(doc.tables)}
            )
        except Exception as e:
            return DocumentExtractionResult(
                text=f"[Error reading DOCX: {str(e)}]",
                page_count=1,
                unclear_sections=[f"Could not read DOCX file: {str(e)}"],
                metadata={"error": str(e)}
            )

    @staticmethod
    def extract_from_image(file_path: str) -> DocumentExtractionResult:
        """Extract text from image files (JPG, PNG, JPEG) using OCR with fallback."""
        unclear_sections = []
        ocr_text = ""
        try:
            image = Image.open(file_path)
            width, height = image.size
            
            # Attempt OCR with pytesseract if tesseract executable is available
            try:
                import pytesseract
                ocr_text = pytesseract.image_to_string(image)
            except Exception as ocr_err:
                unclear_sections.append(f"OCR engine notice: Tesseract engine not running or image low resolution ({str(ocr_err)}).")
                ocr_text = ""

            ocr_text = ocr_text.strip()
            if not ocr_text:
                ocr_text = f"[Image Document: {os.path.basename(file_path)} ({width}x{height}px). OCR text was sparse or low-contrast.]"
                unclear_sections.append("Image resolution or contrast may have affected OCR extraction quality.")

            return DocumentExtractionResult(
                text=FileExtractor.clean_extracted_text(ocr_text),
                page_count=1,
                unclear_sections=unclear_sections,
                metadata={"format": "Image", "dimensions": f"{width}x{height}"}
            )
        except Exception as e:
            return DocumentExtractionResult(
                text=f"[Error opening image: {str(e)}]",
                page_count=1,
                unclear_sections=[f"Failed to process image: {str(e)}"],
                metadata={"error": str(e)}
            )

    @staticmethod
    def extract(file_path: str, original_filename: str) -> DocumentExtractionResult:
        ext = os.path.splitext(original_filename)[1].lower()
        if ext == ".pdf":
            return FileExtractor.extract_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return FileExtractor.extract_from_docx(file_path)
        elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp"]:
            return FileExtractor.extract_from_image(file_path)
        else:
            # Fallback read text
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                return DocumentExtractionResult(
                    text=FileExtractor.clean_extracted_text(content),
                    page_count=1,
                    unclear_sections=[],
                    metadata={"format": "Text"}
                )
            except Exception as e:
                return DocumentExtractionResult(
                    text=f"[Unsupported file format: {ext}]",
                    page_count=1,
                    unclear_sections=[f"Unsupported file extension: {ext}"],
                    metadata={"error": str(e)}
                )

    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """Clean extra spaces, normalize line endings, and preserve tabular structure."""
        if not text:
            return ""
        # Normalize CRLF
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Collapse >3 consecutive newlines
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Strip trailing whitespaces per line
        lines = [line.rstrip() for line in text.split("\n")]
        return "\n".join(lines).strip()
